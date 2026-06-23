import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

doc = Document()

doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading(text, level):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table(data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in data:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
    return table

add_heading('好利航国际物流管理系统', level=0)
doc.add_paragraph('详细需求说明书').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('文档版本：V1.0').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('创建日期：2026年5月').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

add_heading('目录', level=1)
doc.add_paragraph('1. 系统概述')
doc.add_paragraph('2. 菜单结构总览')
doc.add_paragraph('3. 角色与权限体系')
doc.add_paragraph('4. 功能模块详细需求')
doc.add_paragraph('   4.1 工作台')
doc.add_paragraph('   4.2 基础资料')
doc.add_paragraph('   4.3 客户管理(CRM)')
doc.add_paragraph('   4.4 运单管理')
doc.add_paragraph('   4.5 提单管理')
doc.add_paragraph('   4.6 仓库操作管理')
doc.add_paragraph('   4.7 产品配置')
doc.add_paragraph('   4.8 业务配置')
doc.add_paragraph('   4.9 财务管理')
doc.add_paragraph('   4.10 客服管理')
doc.add_paragraph('   4.11 报表管理')
doc.add_paragraph('   4.12 待办')
doc.add_paragraph('   4.13 权限管理')
doc.add_paragraph('   4.14 物料管理')
doc.add_paragraph('   4.15 外部对接接口')
doc.add_paragraph('   4.16 移动端APP')
doc.add_paragraph('   4.17 系统配置')
doc.add_paragraph('5. 界面设计规范')
doc.add_paragraph('6. 弹窗交互标准')
doc.add_paragraph('7. 验收清单')
doc.add_page_break()

add_heading('1. 系统概述', level=1)
add_paragraph('好利航国际物流管理系统是一套面向国际物流企业的综合业务管理平台，涵盖客户管理、运单管理、仓库管理、财务管理等核心业务模块，支持多角色权限管理和多语言切换。')

add_heading('2. 菜单结构总览', level=1)
menu_structure = [
    ['序号', '一级菜单', '二级菜单', '功能描述'],
    ['1', '工作台', '数据看板', '展示关键业务指标和待办事项'],
    ['2', '基础资料', '服务商管理', '管理物流服务商信息'],
    ['3', '基础资料', '员工管理', '管理企业员工信息'],
    ['4', '基础资料', '银行账户', '管理企业银行账户'],
    ['5', '客户管理(CRM)', '客户管理', '管理客户信息和状态'],
    ['6', '客户管理(CRM)', '公海池管理', '管理公海客户资源'],
    ['7', '运单管理', '下单录入', '录入运单预报信息'],
    ['8', '运单管理', '整柜下单', '整柜业务运单录入'],
    ['9', '提单管理', '提单录入', '提单信息录入'],
    ['10', '仓库操作管理', '入仓操作', '仓库入仓业务处理'],
    ['11', '产品配置', '产品管理', '管理物流产品'],
    ['12', '业务配置', '风控规则', '配置业务风控规则'],
    ['13', '财务管理', '应收管理', '管理应收账款'],
    ['14', '客服管理', '工单管理', '处理客户工单'],
    ['15', '报表管理', '业务报表', '生成业务统计报表'],
    ['16', '权限管理', '角色管理', '配置角色权限'],
    ['17', '系统配置', '系统设置', '系统参数配置']
]
add_table(menu_structure[1:], menu_structure[0])

add_heading('3. 角色与权限体系', level=1)
add_paragraph('系统支持4种角色：超级管理员、管理员、业务员、客服。')

add_heading('4. 功能模块详细需求', level=1)

add_heading('4.1 工作台', level=2)
add_paragraph('工作台是用户登录系统后的首页，展示关键业务指标和待办事项。')
add_paragraph('功能特点：')
doc.add_paragraph('• 数据看板：展示今日运单数、待审核数、异常告警等关键指标')
doc.add_paragraph('• 快捷入口：提供常用功能的快速访问')
doc.add_paragraph('• 待办事项：显示需要处理的任务列表')

add_heading('4.2 基础资料', level=2)

add_heading('4.2.1 服务商管理', level=3)
add_paragraph('菜单路径：基础资料 → 服务商管理')
add_paragraph('页面类型：列表管理页')
add_paragraph('列表字段：')
service_provider_fields = [
    ['字段名称', '字段说明', '必填', '数据类型', '字段长度/格式'],
    ['服务商编码', '服务商唯一标识，系统自动生成', '是', '字符串', '固定10位，SERV+6位数字'],
    ['服务商名称', '服务商公司全称', '是', '字符串', '最大100字符'],
    ['服务类型', '服务类型分类：专线/整柜/空运/快递/铁路', '是', '枚举', '下拉选择'],
    ['联系人', '服务商对接人姓名', '否', '字符串', '最大50字符'],
    ['联系电话', '服务商联系电话', '否', '字符串', '手机号或固定电话格式'],
    ['结算方式', '结算类型：月结/周结/日结/预付', '是', '枚举', '下拉选择'],
    ['状态', '服务商状态：启用/停用', '是', '枚举', '标签展示']
]
add_table(service_provider_fields[1:], service_provider_fields[0])

add_paragraph('功能按钮：新增、编辑、查看、启用/停用、导出')

add_paragraph('新增/编辑弹窗说明：')
doc.add_paragraph('• 弹窗类型：右侧滑入弹窗，宽度70%')
doc.add_paragraph('• 弹窗字段：服务商编码(自动生成)、服务商名称(必填)、服务类型(下拉)、联系人、联系电话、邮箱地址、合作渠道、结算方式、银行账户、备注')
doc.add_paragraph('• 底部按钮：取消、保存')

add_heading('4.2.2 员工管理', level=3)
add_paragraph('菜单路径：基础资料 → 员工管理')
add_paragraph('页面类型：列表管理页')
add_paragraph('功能按钮：新增、编辑、查看、重置密码、启用/停用、导出')

add_heading('4.2.3 银行账户', level=3)
add_paragraph('菜单路径：基础资料 → 银行账户')
add_paragraph('页面类型：列表管理页')

add_heading('4.3 客户管理(CRM)', level=2)

add_heading('4.3.1 客户管理', level=3)
add_paragraph('菜单路径：客户管理(CRM) → 客户管理')
add_paragraph('页面类型：列表管理页')
customer_fields = [
    ['字段名称', '字段说明', '必填', '数据类型', '字段长度/格式'],
    ['客户代码', '客户唯一标识，系统自动生成', '是', '字符串', '固定10位，CUST+6位数字'],
    ['客户简称', '客户简称，用于日常业务称呼', '是', '字符串', '最大50字符'],
    ['客户全称', '客户公司全称，用于合同和发票', '是', '字符串', '最大200字符'],
    ['客户类型', '客户类型：货主/同行/同行直客', '是', '枚举', '下拉选择'],
    ['客户等级', '客户等级分类：A/B/C/D', '否', '枚举', '下拉选择'],
    ['联系人', '主要联系人姓名', '否', '字符串', '最大50字符'],
    ['客户邮箱', '客户官方邮箱地址', '否', '字符串', '邮箱格式']
]
add_table(customer_fields[1:], customer_fields[0])

add_paragraph('功能按钮：新增、编辑、查看、分配销售、转入公海、启用/停用、导出')

add_paragraph('客户分配弹窗说明：')
doc.add_paragraph('• 弹窗类型：居中弹窗，宽度500px')
doc.add_paragraph('• 弹窗字段：客户信息(只读)、当前负责人(只读)、新负责人(必填)、分配原因、通知方式')
doc.add_paragraph('• 底部按钮：取消、确认分配')

add_paragraph('新增/编辑客户弹窗说明：')
doc.add_paragraph('• 弹窗类型：右侧滑入弹窗，宽度70%')
doc.add_paragraph('• 弹窗字段：客户代码(自动生成)、客户简称(必填)、客户全称(必填)、客户类型(必填)、客户等级、联系人、联系电话、客户邮箱、客户来源、结算周期、信用额度、所属业务员、备注')
doc.add_paragraph('• 底部按钮：取消、保存')

add_heading('4.3.2 公海池管理', level=3)
add_paragraph('菜单路径：客户管理(CRM) → 公海池管理')
add_paragraph('页面类型：列表管理页')
add_paragraph('功能按钮：领取、分配、查看')

add_paragraph('领取客户弹窗说明：')
doc.add_paragraph('• 弹窗类型：居中弹窗，宽度450px')
doc.add_paragraph('• 弹窗字段：客户信息(只读)、进入公海时间(只读)、公海原因(只读)、领取原因(必填)')
doc.add_paragraph('• 底部按钮：取消、确认领取')

add_heading('4.4 运单管理', level=2)

add_heading('4.4.1 下单录入', level=3)
add_paragraph('菜单路径：运单管理 → 下单录入')
add_paragraph('页面类型：录入型页面')
add_paragraph('功能按钮：提交预报、保存草稿、附件上传、重置、需求说明')

add_paragraph('提交预报弹窗说明：')
doc.add_paragraph('• 弹窗类型：居中弹窗，宽度600px')
doc.add_paragraph('• 弹窗字段：运单信息(只读)、审核节点(必填)、提交备注、通知方式')
doc.add_paragraph('• 底部按钮：取消、确认提交')

add_paragraph('附件上传弹窗说明：')
doc.add_paragraph('• 弹窗类型：居中弹窗，宽度550px')
doc.add_paragraph('• 支持文件类型：PDF、图片(JPG/PNG)、Excel、Word')
doc.add_paragraph('• 文件大小限制：单个文件不超过10MB')
doc.add_paragraph('• 底部按钮：取消、确定')

add_heading('4.4.2 整柜下单', level=3)
add_paragraph('菜单路径：运单管理 → 整柜下单')
add_paragraph('页面类型：列表管理页')

add_heading('4.5 提单管理', level=2)

add_heading('4.5.1 提单录入', level=3)
add_paragraph('菜单路径：提单管理 → 提单录入')

add_heading('4.5.2 提单列表', level=3)
add_paragraph('菜单路径：提单管理 → 提单列表')

add_heading('4.5.3 提单审核', level=3)
add_paragraph('菜单路径：提单管理 → 提单审核')
add_paragraph('页面类型：列表管理页')

add_paragraph('提单审核弹窗说明：')
doc.add_paragraph('• 弹窗类型：居中弹窗，宽度600px')
doc.add_paragraph('• 弹窗字段：提单详情(只读)、审核类型(只读)、提交人信息(只读)、审核意见(必填)、审核结果(必填)')
doc.add_paragraph('• 底部按钮：取消、审核驳回、审核通过')

add_heading('4.5.4 提单修改', level=3)
add_paragraph('菜单路径：提单管理 → 提单修改')
add_paragraph('页面类型：列表管理页')

add_heading('4.6 仓库操作管理', level=2)

add_heading('4.6.1 入仓操作', level=3)
add_paragraph('菜单路径：仓库操作管理-国内 → 入仓操作')
add_paragraph('页面类型：录入型页面')
add_paragraph('功能按钮：保存入仓、打印标签、异常登记、重置、需求说明、刷新数据')

add_paragraph('异常登记弹窗说明：')
doc.add_paragraph('• 弹窗类型：居中弹窗，宽度550px')
doc.add_paragraph('• 弹窗字段：入仓单号(只读)、异常类型(必填)、责任方(必填)、异常等级(必填)、异常说明(必填)、照片上传、处理建议')
doc.add_paragraph('• 底部按钮：取消、确认登记')

add_heading('4.6.2 理货', level=3)
add_paragraph('菜单路径：仓库操作管理-国内 → 理货')

add_heading('4.6.3 上托', level=3)
add_paragraph('菜单路径：仓库操作管理-国内 → 上托')

add_heading('4.6.4 出库', level=3)
add_paragraph('菜单路径：仓库操作管理-国内 → 出库')

add_heading('4.7 产品配置', level=2)

add_heading('4.7.1 产品管理', level=3)
add_paragraph('菜单路径：产品配置 → 产品管理')

add_heading('4.7.4 询价报价', level=3)
add_paragraph('菜单路径：产品配置 → 询价报价')
add_paragraph('页面类型：录入型页面')
add_paragraph('功能按钮：确认提交、保存草稿、重置、需求说明')

add_paragraph('报价确认弹窗说明：')
doc.add_paragraph('• 弹窗类型：居中弹窗，宽度600px')
doc.add_paragraph('• 弹窗字段：客户信息(只读)、货物信息(只读)、报价明细(只读)、报价有效期(必填)、备注')
doc.add_paragraph('• 底部按钮：取消、确认报价')

add_heading('4.8 业务配置', level=2)

add_heading('4.8.1 风控规则', level=3)
add_paragraph('菜单路径：业务配置 → 风控规则')

add_heading('4.9 财务管理', level=2)

add_heading('4.9.1 客户与账期', level=3)
add_paragraph('菜单路径：财务管理 → 客户与账期')

add_heading('4.9.2 应收管理', level=3)
add_paragraph('菜单路径：财务管理 → 应收管理')

add_heading('4.9.3 应付与成本', level=3)
add_paragraph('菜单路径：财务管理 → 应付与成本')

add_heading('4.9.4 汇率与开票', level=3)
add_paragraph('菜单路径：财务管理 → 汇率与开票')

add_heading('4.10 客服管理', level=2)

add_heading('4.10.1 工单管理', level=3)
add_paragraph('菜单路径：客服管理 → 工单管理')

add_heading('4.10.2 投诉管理', level=3)
add_paragraph('菜单路径：客服管理 → 投诉管理')

add_heading('4.11 报表管理', level=2)

add_heading('4.11.1 业务报表', level=3)
add_paragraph('菜单路径：报表管理 → 业务报表')

add_heading('4.11.2 财务报表', level=3)
add_paragraph('菜单路径：报表管理 → 财务报表')

add_heading('4.12 待办', level=2)

add_heading('4.12.1 待办中心', level=3)
add_paragraph('菜单路径：待办 → 待办中心')

add_heading('4.13 权限管理', level=2)

add_heading('4.13.1 用户管理', level=3)
add_paragraph('菜单路径：权限管理 → 用户管理')

add_heading('4.13.2 角色管理', level=3)
add_paragraph('菜单路径：权限管理 → 角色管理')
add_paragraph('页面类型：列表管理页')

add_paragraph('角色权限配置弹窗说明：')
doc.add_paragraph('• 弹窗类型：右侧滑入弹窗，宽度70%')
doc.add_paragraph('• 左侧区域：角色名称(必填)、角色描述、状态、创建时间(只读)、修改时间(只读)')
doc.add_paragraph('• 右侧区域三列：菜单权限(树形结构)、字段权限(表格形式)、查询条件权限')
doc.add_paragraph('• 底部按钮：取消、保存角色')

add_heading('4.13.3 菜单管理', level=3)
add_paragraph('菜单路径：权限管理 → 菜单管理')

add_heading('4.13.4 日志查询', level=3)
add_paragraph('菜单路径：权限管理 → 日志查询')

add_heading('4.14 物料管理', level=2)

add_heading('4.14.1 物料列表', level=3)
add_paragraph('菜单路径：物料管理 → 物料列表')

add_heading('4.15 外部对接接口', level=2)

add_heading('4.15.1 接口配置', level=3)
add_paragraph('菜单路径：外部对接接口 → 接口配置')

add_heading('4.16 移动端APP', level=2)

add_heading('4.16.1 APP功能', level=3)
add_paragraph('菜单路径：移动端APP → APP功能')

add_heading('4.17 系统配置', level=2)

add_heading('4.17.1 系统设置', level=3)
add_paragraph('菜单路径：系统配置 → 系统设置')

add_heading('5. 界面设计规范', level=1)
add_paragraph('• 设计原则：简洁、高效、统一')
add_paragraph('• 色彩规范：主色调蓝色(#1890ff)，辅助色灰色系')
add_paragraph('• 字体规范：微软雅黑，14px标准字号')
add_paragraph('• 间距规范：标准间距16px，组件间距8px')
add_paragraph('• 按钮规范：主按钮蓝色，次按钮灰色边框，文字按钮无背景')

add_heading('6. 弹窗交互标准', level=1)
add_paragraph('• 弹窗类型：右侧滑入弹窗(70%宽度)、居中弹窗(固定宽度)')
add_paragraph('• 滑入动画：从右向左滑入，带过渡效果')
add_paragraph('• 关闭方式：点击关闭按钮、点击遮罩层、ESC键')
add_paragraph('• 按钮布局：取消按钮在左，主操作按钮在右')

add_heading('7. 验收清单', level=1)
add_paragraph('• 功能验收：所有菜单功能正常访问和操作')
add_paragraph('• UI验收：界面风格统一，响应式布局正常')
add_paragraph('• 交互验收：弹窗动画流畅，按钮点击响应正常')
add_paragraph('• 权限验收：不同角色菜单过滤正确')
add_paragraph('• 多语言验收：中/英/法三种语言切换正常')

doc.save('好利航国际物流_详细需求说明书.docx')
print('Word文档生成成功！')